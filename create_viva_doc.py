import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = docx.Document()

# Title
title = doc.add_heading('Viva Voce Preparation Guide', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph("This document contains extremely simple explanations to help you prepare for your final viva voce. Imagine you are explaining these concepts to someone who knows nothing about computers or AI.\n")

# 1. SHAP
doc.add_heading('1. What is SHAP? (And why did we use it?)', level=1)
doc.add_paragraph("Imagine our AI model is a detective who looks at a suspect (a review) and says, \"This guy is guilty (fake review)!\" But the judge (you) asks, \"Why? How do you know?\"")
doc.add_paragraph("Normally, AI is a 'black box'—it gives an answer but can't explain why. SHAP is the translator that forces the AI to explain its math in human terms. It stands for 'SHapley Additive exPlanations'.")
doc.add_paragraph("What SHAP does:")
doc.add_paragraph("• It breaks down the AI's final decision into tiny pieces.", style='List Bullet')
doc.add_paragraph("• It tells us EXACTLY which clues the AI used. For example, it will say: \"I think this review is 80% fake. The fact that the user posted 50 reviews in one day added +40% to my suspicion. The fact that they only give 5-star ratings added +30%.\"", style='List Bullet')
doc.add_paragraph("Why we used it in our project:")
doc.add_paragraph("In our dissertation, we didn't just want to build an AI that catches fake reviews. We wanted to PROVE that 'User Behavior' is the biggest giveaway. By using SHAP, we created a dashboard where anyone can click on a review and physically see that the AI caught the spammer not because of their words, but because of their suspicious habits (like posting too fast). SHAP makes our AI trustworthy.")

# 2. Behavioral Features Glossary
doc.add_heading('2. Behavioral Features Glossary (Simple Terms)', level=1)
doc.add_paragraph("In our project, we didn't just read the text of the reviews. We tracked the 'habits' of the users and the 'patterns' of the restaurants. We created 55 of these habits (features). Here are the main ones explained simply:")

doc.add_heading('Reviewer Habits (Catching the Spammer)', level=2)
doc.add_paragraph("• Burstiness (max_reviews_per_day): Is this person posting 50 reviews in a single day? Normal humans don't do that. Spambots do.")
doc.add_paragraph("• Extreme Rating Ratio (EXRR): Does this person only ever give 1-star or 5-star ratings? Real people give 3 or 4 stars sometimes. Spammers are paid to either destroy or boost a rating.")
doc.add_paragraph("• Similarity Score (MCS): Does this person copy and paste the exact same review text for 10 different restaurants? ")
doc.add_paragraph("• Review Frequency (review_count): How many reviews has this person written in total? ")
doc.add_paragraph("• Early Review Ratio (AFPPR): Is this person always the very first person to review a newly opened restaurant? Spammers do this to artificially boost a place right when it opens.")

doc.add_heading('Product Patterns (Catching the Target)', level=2)
doc.add_paragraph("• Activity Spike (product_review_velocity): Did this restaurant suddenly get 500 reviews in one weekend when they normally get 2 a month? That's a huge red flag.")
doc.add_paragraph("• Rating Deviation (product_rating_std): Is the restaurant's star rating violently bouncing up and down over time? This shows two groups might be fighting (fake positive reviews vs real angry customers).")
doc.add_paragraph("• Repeated Positive Ratio (RPR): Are 99% of the reviews for this restaurant overwhelmingly positive, with no natural complaints? ")

# 3. Literature Survey
doc.add_heading('3. Literature Survey (The History of Fake Review Detection)', level=1)
doc.add_paragraph("When researchers first started trying to catch fake reviews, they went through three main phases. Here is the simple story:")

doc.add_heading('Phase 1: Just reading the words (Text-Based)', level=2)
doc.add_paragraph("In the beginning, scientists tried to catch fake reviews by looking at the text (using Natural Language Processing). They looked for spelling mistakes, overly emotional words, or weird grammar. \nProblem: Spammers got smart. They started writing perfect, realistic-sounding reviews. So, text alone stopped working.")

doc.add_heading('Phase 2: Watching what they do (Behavior-Based)', level=2)
doc.add_paragraph("Researchers realized that even if a spammer writes a perfect review, their *actions* give them away. A famous paper by Rayana and Akoglu (2015) proved that looking at a user's network and habits—like posting 100 reviews in an hour—is the best way to catch them. ")

doc.add_heading('Phase 3: The Hybrid Approach (What we did)', level=2)
doc.add_paragraph("Today, the best approach is to combine both! In our dissertation, we compared three AI models:")
doc.add_paragraph("• Model 1: Used messy, raw data.")
doc.add_paragraph("• Model 2: Used clean data, but only looked at the Text.")
doc.add_paragraph("• Model 3: Used clean data, and looked at BOTH Text and User Behavior.")
doc.add_paragraph("Our literature survey proved that to beat modern spammers, you must track their behavior. And our final experiment (Model 3) successfully proved this theory, showing a massive jump in accuracy when we added our 55 behavioral features.")

# Save
doc.save('Viva_Preparation_Guide.docx')
print("Successfully created Viva_Preparation_Guide.docx")
