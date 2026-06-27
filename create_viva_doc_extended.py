import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = docx.Document()

# Title
title = doc.add_heading('Viva Voce Preparation Guide (Extended)', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph("This document contains elaborated yet simple explanations to help you prepare for your final viva voce. It is designed so you can explain these complex concepts to someone without any coding or AI background.\n")

# -----------------------------------------------------------------------------
# 1. ELABORATED SHAP EXPLANATION
# -----------------------------------------------------------------------------
doc.add_heading('1. What is SHAP? (Elaborated but Simple)', level=1)

doc.add_heading('The Problem: AI is a "Black Box"', level=2)
doc.add_paragraph("Normally, Machine Learning models are like a magic 'black box'. You feed them data (a restaurant review), and they spit out an answer (\"This review is Fake!\"). If you ask the AI, \"Why did you think it's fake?\" the AI cannot answer. It just says, \"Because the math told me so.\" In the real world, especially when detecting fraud, you cannot just trust a black box. You need proof.")

doc.add_heading('The Solution: SHAP', level=2)
doc.add_paragraph("SHAP stands for 'SHapley Additive exPlanations'. It is a tool that forces the AI to open its brain and show us exactly how it made its decision. It acts like a translator between the complex math of the AI and human understanding.")

doc.add_heading('How it works (The Game Theory Analogy)', level=2)
doc.add_paragraph("To explain SHAP to the examiners, use this simple analogy:")
doc.add_paragraph("Imagine a team of construction workers who built a house together. When the house is finished, they are paid $100,000. But how do you divide the money fairly? Some workers laid the bricks, some painted the walls, and some just swept the floor. You have to calculate EXACTLY how much value each specific worker added to the final house.")
doc.add_paragraph("SHAP does exactly this for AI. Instead of 'workers building a house', we have 'data clues making a prediction'.")
doc.add_paragraph("• The 'House' is the AI's final decision (e.g., \"I am 90% sure this review is fake\").")
doc.add_paragraph("• The 'Workers' are the clues (e.g., the user's burstiness, their extreme ratings, the text they wrote).")
doc.add_paragraph("SHAP calculates exactly how much each clue contributed to the final 90% decision.")

doc.add_heading('How we used it in our project:', level=2)
doc.add_paragraph("In our Streamlit dashboard on the 'SHAP Analysis' page, we use SHAP to explain individual predictions. For example, when you click on a fake review, SHAP might output:")
doc.add_paragraph("• Base suspicion: 10%")
doc.add_paragraph("• Fact: User posted 50 reviews in one day (Added +50% suspicion)")
doc.add_paragraph("• Fact: User only gives 5-star ratings (Added +20% suspicion)")
doc.add_paragraph("• Fact: The review text was perfectly normal (Subtracted -5% suspicion)")
doc.add_paragraph("• Final AI Decision: 75% Fake.")
doc.add_paragraph("This proves to the examiners that our AI isn't just randomly guessing. It proves that our 'Behavioral Features' are actually working and actively catching spammers.")

# -----------------------------------------------------------------------------
# 2. ALL 55 BEHAVIORAL FEATURES
# -----------------------------------------------------------------------------
doc.add_heading('2. All 55 Behavioral Features Glossary', level=1)
doc.add_paragraph("We engineered 55 unique mathematical features that track the 'habits' of users and the 'patterns' of restaurants. These are the clues the AI uses to catch spammers. Here is the complete list explained simply:")

doc.add_heading('Part A: Reviewer-Level Features (Catching the Spammer)', level=2)
doc.add_paragraph("These 28 features analyze the person writing the review. (They end with '_r' for Reviewer).")

reviewer_features = [
    ("1. ARD_r (Average Rating Deviation)", "How much this user's ratings typically differ from the crowd's average rating."),
    ("2. WRD_r (Weighted Rating Deviation)", "Similar to ARD, but weighted by the number of reviews the product has."),
    ("3. MRD_r (Maximum Rating Deviation)", "The single most extreme difference between this user's rating and the crowd's rating."),
    ("4. max_reviews_per_day_r (Burstiness)", "The highest number of reviews this user has ever posted in a single 24-hour window."),
    ("5. review_interval_cv_r (Posting Irregularity)", "Detects if a user posts naturally over time or in sudden, robotic bursts."),
    ("6. ERR_r (Early Reviewer Ratio)", "How frequently this user writes reviews during the very early days of a product's launch."),
    ("7. RPR_r (Repeated Positive Ratio)", "The proportion of this user's reviews that are overwhelmingly positive."),
    ("8. RNR_r (Repeated Negative Ratio)", "The proportion of this user's reviews that are overwhelmingly negative."),
    ("9. EXRR_r (Extreme Rating Ratio)", "The percentage of the user's reviews that are either 1-star or 5-stars (highly polarized)."),
    ("10. TRRR_r (Top Rating Ratio)", "The percentage of the user's reviews that are exactly 5-stars."),
    ("11. BRRR_r (Bottom Rating Ratio)", "The percentage of the user's reviews that are exactly 1-star."),
    ("12. MCS_r (Max Content Similarity)", "The highest text similarity score between any two reviews written by this user (catches copy-pasting)."),
    ("13. ACS_r (Average Content Similarity)", "How repetitive or templated the user's writing style is across all their reviews."),
    ("14. AFPPR_r (First Post Pattern Ratio)", "The frequency at which the user is among the first few reviewers of a product."),
    ("15. ASPPR_r (Second Post Pattern Ratio)", "The frequency at which the user follows immediately after the first reviewer."),
    ("16. AFTAPP_r (Average Time to Post)", "The average speed at which the user posts reviews after creating their account."),
    ("17. vader_compound_mean_r (Avg Emotion)", "The average emotional positivity/negativity of their text (calculated via VADER NLP)."),
    ("18. vader_compound_std_r (Emotion Volatility)", "How wildly the user's emotional tone swings from review to review."),
    ("19. vader_extreme_ratio_r (Extreme Emotion)", "The percentage of reviews that use intensely emotional or dramatic language."),
    ("20. avg_word_count_r (Avg Word Count)", "The average number of words this user writes per review."),
    ("21. reviewer_avg_rating_r (Historical Avg Rating)", "The mean star rating this user gives out across all time."),
    ("22. rating_entropy_r (Rating Unpredictability)", "The mathematical entropy (unpredictability) of the user's rating distribution."),
    ("23. review_count_r (Total Review Count)", "The absolute total number of reviews this user has published."),
    ("24. FRR_r (First Review Ratio)", "The percentage of times this user was the very first person to review a business."),
    ("25. pct_single_review_products_r (Obscure Product Ratio)", "The percentage of reviews left on businesses that have exactly 1 review total."),
    ("26. weekend_ratio_r (Weekend Activity Ratio)", "The proportion of reviews published on Saturdays or Sundays."),
    ("27. unique_products_ratio_r (Unique Product Ratio)", "Measures if the user reviews different places or repeatedly reviews the same place."),
    ("28. MNR_r (Max Reviews per Product)", "The highest number of times this user has reviewed a single specific business.")
]

table1 = doc.add_table(rows=1, cols=2)
table1.style = 'Table Grid'
hdr_cells1 = table1.rows[0].cells
hdr_cells1[0].text = 'Feature'
hdr_cells1[1].text = 'Simple Explanation'
for feature, expl in reviewer_features:
    row_cells = table1.add_row().cells
    row_cells[0].text = feature
    row_cells[1].text = expl

doc.add_paragraph("\n")
doc.add_heading('Part B: Product-Level Features (Catching the Target)', level=2)
doc.add_paragraph("These 27 features analyze the restaurant/business receiving the review to detect suspicious campaigns directed at them. (They end with '_p' for Product).")

product_features = [
    ("29. ARD_p (Average Rating Deviation)", "The average deviation of incoming ratings from the product's historical mean."),
    ("30. WRD_p (Weighted Rating Deviation)", "Similar to ARD, weighted by reviewer credibility or review volume."),
    ("31. MRD_p (Maximum Rating Deviation)", "The largest spike or drop in ratings compared to the norm."),
    ("32. max_reviews_per_day_p (Activity Spike)", "The maximum number of reviews the business received in a single 24-hour window."),
    ("33. review_interval_cv_p (Pace Irregularity)", "Measures if reviews trickle in naturally or arrive in suspicious, coordinated clumps."),
    ("34. ERR_p (Early Review Ratio)", "The concentration of reviews received very early in the business's listing lifecycle."),
    ("35. RPR_p (Repeated Positive Ratio)", "The overall density of purely positive reviews on the product."),
    ("36. RNR_p (Repeated Negative Ratio)", "The overall density of purely negative reviews (indicative of review-bombing)."),
    ("37. EXRR_p (Extreme Rating Ratio)", "The percentage of reviews that are either 1-star or 5-stars."),
    ("38. TRRR_p (Top Rating Ratio)", "The percentage of all incoming reviews that are exactly 5-stars."),
    ("39. BRRR_p (Bottom Rating Ratio)", "The percentage of all incoming reviews that are exactly 1-star."),
    ("40. MCS_p (Max Content Similarity)", "The highest similarity score between any two reviews on this product (spots bot swarms)."),
    ("41. ACS_p (Average Content Similarity)", "How repetitive or templated the reviews for this business are overall."),
    ("42. AFPPR_p (Early Reviewer Concentration)", "The ratio of reviews coming from accounts that specialize in reviewing products early."),
    ("43. ASPPR_p (Secondary Reviewer Concentration)", "The ratio of reviews coming from accounts that immediately follow early reviewers."),
    ("44. AFTAPP_p (Average Review Interval)", "The average time elapsed between consecutive reviews for this business."),
    ("45. vader_compound_mean_p (Product Sentiment)", "The average emotional tone of all text written about this business."),
    ("46. vader_compound_std_p (Sentiment Volatility)", "How much the emotional tone fluctuates in the product's reviews."),
    ("47. vader_extreme_ratio_p (Extreme Emotion)", "The proportion of reviews on this product using intensely emotional language."),
    ("48. avg_word_count_p (Product Review Length)", "The average length of reviews left on this business."),
    ("49. reviewer_avg_rating_p (Reviewer Baseline)", "The average historical rating of the people currently reviewing this business."),
    ("50. rating_entropy_p (Rating Unpredictability)", "The mathematical entropy of the product's overall star rating distribution."),
    ("51. review_count_p (Total Product Reviews)", "The absolute total number of reviews this business has accumulated."),
    ("52. MNR_p (Max Reviews by Single User)", "The most reviews left by a single individual on this specific business."),
    ("53. product_review_velocity (Momentum)", "The rate of acceleration at which the business is gaining new reviews."),
    ("54. product_rating_std (Rating Volatility)", "The standard deviation of the business's star rating over time (detects manipulation swings)."),
    ("55. product_fake_ratio (Historical Fake Ratio)", "The historical proportion of reviews on this product that have been flagged as fake previously.")
]

table2 = doc.add_table(rows=1, cols=2)
table2.style = 'Table Grid'
hdr_cells2 = table2.rows[0].cells
hdr_cells2[0].text = 'Feature'
hdr_cells2[1].text = 'Simple Explanation'
for feature, expl in product_features:
    row_cells = table2.add_row().cells
    row_cells[0].text = feature
    row_cells[1].text = expl

doc.save('Viva_Preparation_Guide_Extended.docx')
print("Successfully created Viva_Preparation_Guide_Extended.docx")
